#!/usr/bin/env python3
"""读取 DOCX 文件并输出文本内容 - 支持通配符和批量处理"""

import sys
import glob
import os
from docx import Document

def read_docx(path, verbose=False):
    """读取 docx 文件，支持通配符"""
    # 用通配符找文件（避免中文文件名编码问题）
    files = glob.glob(path)
    if not files and not path.endswith('.docx'):
        files = glob.glob(path + "*.docx")
    
    if not files:
        print(f"❌ 找不到文件：{path}", file=sys.stderr)
        sys.exit(1)
    
    if verbose:
        print(f"📁 找到 {len(files)} 个文件\n")
    
    success = 0
    failed = 0
    
    for f in sorted(files):
        try:
            doc = Document(f)
            # 最快提取方式：用 generator 避免中间 list
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            
            # 如果还想更快，可以直接访问底层元素（极致场景）
            # text = '\n'.join(run.text for para in doc.paragraphs for run in para.runs if run.text.strip())
            
            print(f"--- {os.path.basename(f)} ---")
            print(text)
            print("\n" + "="*60 + "\n")
            success += 1
        except Exception as e:
            print(f"处理 {f} 失败：{e}", file=sys.stderr)
            failed += 1
    
    if verbose:
        print(f"✅ 成功：{success} | ❌ 失败：{failed}")
    
    return success, failed

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python3 read-docx.py <文件路径> [--verbose]", file=sys.stderr)
        print("示例：", file=sys.stderr)
        print("  python3 read-docx.py exams/HDMI*.docx", file=sys.stderr)
        print("  python3 read-docx.py /path/to/*.docx -v", file=sys.stderr)
        sys.exit(1)
    
    verbose = '-v' in sys.argv or '--verbose' in sys.argv
    path = sys.argv[1]
    if path in ['-v', '--verbose']:
        print("错误：请提供文件路径", file=sys.stderr)
        sys.exit(1)
    
    read_docx(path, verbose)
