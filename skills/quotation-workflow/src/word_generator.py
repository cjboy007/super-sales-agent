#!/usr/bin/env python3
"""
报价单 Word 生成器
"""

from dataclasses import dataclass
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


@dataclass
class GenerationResult:
    """生成结果"""
    success: bool
    output_path: Optional[str] = None
    error: Optional[str] = None


def generate_quotation_word(data: dict, output_path: str) -> GenerationResult:
    """
    生成报价单 Word 文档
    
    Args:
        data: 报价单数据（包含 customer, products, quotation, trade_terms）
        output_path: 输出文件路径
    
    Returns:
        GenerationResult: 生成结果
    """
    try:
        # 创建 Document
        doc = Document()
        
        # 设置样式 - Times New Roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # 标题
        title = doc.add_heading('QUOTATION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 公司信息
        company = doc.add_paragraph('FARREACH ELECTRONIC CO LIMITED')
        company.runs[0].bold = True
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 报价单信息
        quotation = data.get('quotation', {})
        doc.add_paragraph(f"Quotation No: {quotation.get('quotation_no', 'N/A')}")
        doc.add_paragraph(f"Date: {quotation.get('date', 'N/A')}")
        doc.add_paragraph(f"Valid Until: {quotation.get('valid_until', 'N/A')}")
        
        # 客户信息
        doc.add_heading('To:', level=2)
        customer = data.get('customer', {})
        doc.add_paragraph(customer.get('company_name', ''), style='Intense Quote')
        doc.add_paragraph(customer.get('address', ''))
        doc.add_paragraph(f"Tel: {customer.get('phone', '')}")
        doc.add_paragraph(f"Email: {customer.get('email', '')}")
        
        # 产品表格
        doc.add_heading('Products', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # 表格标题
        headers = ['No.', 'Description', 'Specification', 'Qty', 'Unit Price', 'Total']
        header_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            header_cells[i].text = text
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 表格数据
        products = data.get('products', [])
        for i, product in enumerate(products, start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = product.get('description', '')
            row[2].text = product.get('specification', '')
            row[3].text = str(product.get('quantity', 0))
            row[4].text = f"${product.get('unit_price', 0):.2f}"
            line_total = product.get('quantity', 0) * product.get('unit_price', 0)
            row[5].text = f"${line_total:.2f}"
            
            # 右对齐价格和总额
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 总额
        subtotal = sum(p.get('quantity', 0) * p.get('unit_price', 0) for p in products)
        doc.add_paragraph()
        total_para = doc.add_paragraph()
        total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_run = total_para.add_run(f'TOTAL: ${subtotal:.2f}')
        total_run.bold = True
        
        # 贸易条款
        doc.add_heading('Terms & Conditions', level=2)
        trade_terms = data.get('trade_terms', {})
        terms = [
            f"Payment: T/T 30% deposit, 70% before shipment",
            f"Delivery: {trade_terms.get('delivery', 'N/A')}",
            f"Incoterms: {trade_terms.get('incoterms', 'N/A')}",
            f"Currency: {trade_terms.get('currency', 'USD')}",
            f"This quotation is valid until {quotation.get('valid_until', 'N/A')}"
        ]
        for term in terms:
            doc.add_paragraph(term, style='List Bullet')
        
        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('Thank you for your inquiry. We look forward to your order.')
        footer_run.italic = True
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # 保存文件
        doc.save(output_path)
        
        return GenerationResult(
            success=True,
            output_path=output_path
        )
        
    except Exception as e:
        return GenerationResult(
            success=False,
            error=str(e)
        )
